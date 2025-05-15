import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_markdown_to_docx(markdown_file, docx_file):
    # 读取markdown文件内容
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # 创建一个新的Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 分割内容为行
    lines = markdown_content.split('\n')
    
    # 表格状态跟踪
    in_table = False
    table_data = []
    table_headers = []
    
    # 记录当前的一级标题和二级标题编号
    current_level1_num = 0
    current_level2_nums = {}  # 使用字典跟踪每个一级标题下的二级标题编号
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 处理标题
        if line.startswith('# '):
            # 标题级别1 (文档标题)
            title_text = line[2:]
            title = doc.add_heading(level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.add_run(title_text)
            run.font.name = '宋体'
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色
            run.bold = True
        
        elif line.startswith('## '):
            # 标题级别2 (一级标题 - 中文序号+、)
            title_text = line[3:]
            current_level1_num += 1
            current_level2_nums[current_level1_num] = 0  # 初始化该一级标题下的二级标题编号
            
            chinese_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
            if current_level1_num <= len(chinese_nums):
                formatted_title = f"{chinese_nums[current_level1_num-1]}、{title_text}"
            else:
                formatted_title = f"{current_level1_num}、{title_text}"
            
            heading = doc.add_heading(level=1)
            run = heading.add_run(formatted_title)
            run.font.name = '宋体'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色
            run.bold = True
        
        elif line.startswith('### '):
            # 标题级别3 (二级标题 - 阿拉伯数字)
            title_text = line[4:]
            
            # 增加当前一级标题下的二级标题编号
            current_level2_nums[current_level1_num] += 1
            
            formatted_title = f"{current_level1_num}.{current_level2_nums[current_level1_num]} {title_text}"
            
            heading = doc.add_heading(level=2)
            run = heading.add_run(formatted_title)
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 设置黑色
            run.bold = True
        
        elif line.startswith('#### '):
            # 标题级别4
            title_text = line[5:]
            heading = doc.add_heading(level=3)
            run = heading.add_run(title_text)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 设置黑色
            run.bold = True
        
        # 处理表格
        elif line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                # 解析表头
                table_headers = [cell.strip() for cell in line[1:-1].split('|')]
                # 检查下一行是否是表格分隔符
                if i + 1 < len(lines) and re.match(r'^\|\s*[-:]+\s*\|', lines[i + 1]):
                    i += 1  # 跳过分隔符行
            else:
                # 解析表格数据行
                # 移除首尾的|
                cells = [cell.strip() for cell in line[1:-1].split('|')]
                # 处理单元格中的<br>标签，替换为换行符
                cells = [cell.replace('<br>', '\n') for cell in cells]
                if len(cells) == len(table_headers):
                    table_data.append(cells)
        
        elif in_table and not line.startswith('|'):
            # 表格结束，创建Word表格
            in_table = False
            if table_headers and table_data:
                # 创建表格
                table = doc.add_table(rows=len(table_data) + 1, cols=len(table_headers))
                table.style = 'Table Grid'
                
                # 添加表头
                header_cells = table.rows[0].cells
                for j, header in enumerate(table_headers):
                    header_cells[j].text = header
                    # 设置表头字体
                    for paragraph in header_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = '宋体'
                
                # 添加数据行
                for row_idx, row in enumerate(table_data):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, cell_text in enumerate(row):
                        row_cells[col_idx].text = cell_text
                        # 设置单元格字体
                        for paragraph in row_cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '宋体'
                
                # 重置表格数据
                table_headers = []
                table_data = []
            
            # 处理当前非表格行
            if line:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.name = '宋体'
        
        # 处理普通段落
        elif line and not in_table:
            # 处理列表项
            if line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                for run in p.runs:
                    run.font.name = '宋体'
            elif re.match(r'^\d+\.\s', line):
                p = doc.add_paragraph(line, style='List Number')
                for run in p.runs:
                    run.font.name = '宋体'
            else:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.name = '宋体'
        
        # 处理空行
        elif not line and not in_table:
            if i + 1 < len(lines) and lines[i + 1].strip():
                doc.add_paragraph()
        
        i += 1
    
    # 如果最后还有表格数据未处理，创建表格
    if in_table and table_headers and table_data:
        table = doc.add_table(rows=len(table_data) + 1, cols=len(table_headers))
        table.style = 'Table Grid'
        
        # 添加表头
        header_cells = table.rows[0].cells
        for j, header in enumerate(table_headers):
            header_cells[j].text = header
            # 设置表头字体
            for paragraph in header_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '宋体'
        
        # 添加数据行
        for row_idx, row in enumerate(table_data):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_text in enumerate(row):
                row_cells[col_idx].text = cell_text
                # 设置单元格字体
                for paragraph in row_cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
    
    # 保存文档
    doc.save(docx_file)
    print(f"已将 {markdown_file} 转换为 {docx_file}")

# 创建输出目录
output_dir = os.path.join(os.path.dirname(__file__), 'output')
os.makedirs(output_dir, exist_ok=True)

# 转换需求文档
requirements_md = os.path.join(os.path.dirname(__file__), 'requirements_zh.md')
requirements_docx = os.path.join(output_dir, '档案检索系统需求分析文档.docx')
convert_markdown_to_docx(requirements_md, requirements_docx)
print(f'已将 {requirements_md} 转换为 {requirements_docx}')

# 转换验收文档
acceptance_md = os.path.join(os.path.dirname(__file__), 'acceptance_zh.md')
acceptance_docx = os.path.join(output_dir, '档案检索系统验收文档.docx')
convert_markdown_to_docx(acceptance_md, acceptance_docx)
print(f'已将 {acceptance_md} 转换为 {acceptance_docx}')
